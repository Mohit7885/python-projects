import cv2
import numpy as np
from docx import Document
import tempfile
import pyttsx3
import datetime
import wikipedia
import webbrowser
import os
import pyjokes
import speech_recognition as sr
from PIL import ImageGrab
import time
from ctypes import POINTER, cast
from comtypes import CLSCTX_ALL
from pycaw.pycaw import AudioUtilities, IAudioEndpointVolume
import random
import platform

# Text-to-speech function
def speak(text):
    print(f"Assistant: {text}")
    try:
        engine = pyttsx3.init()
        engine.say(text)
        engine.runAndWait()
    except:
        print("Speech output not supported in this environment.")

# Greet user based on current time
def wish_user():
    hour = int(datetime.datetime.now().hour)
    if hour < 12:
        speak("Good Morning!")
    elif hour < 18:
        speak("Good Afternoon!")
    else:
        speak("Good Evening!")
    speak("Boss, how can I help you?")

# Function to listen for any speech
def listen_for_speech():
    recognizer = sr.Recognizer()
    with sr.Microphone() as source:
        print("Listening for wake word ('Ankit')...")
        recognizer.adjust_for_ambient_noise(source, duration=0.1)
        audio = recognizer.listen(source)
    try:
        query = recognizer.recognize_google(audio)
        print(f"You said: {query}")
        return query.lower()
    except sr.UnknownValueError:
        return ""
    except sr.RequestError:
        speak("Speech recognition service is unavailable.")
        return ""

# Function to listen for a command after wake word
def listen_for_command():
    recognizer = sr.Recognizer()
    with sr.Microphone() as source:
        speak("Yes Boss.")
        recognizer.adjust_for_ambient_noise(source, duration=0.5)
        audio = recognizer.listen(source)
    try:
        command = recognizer.recognize_google(audio)
        print(f"Command received: {command}")
        return command.lower()
    except sr.UnknownValueError:
        speak("Sorry, I didn't catch that.")
        return ""
    except sr.RequestError:
        speak("Speech service is currently unavailable.")
        return ""

# Screenshot function
def take_screenshot():
    try:
        img = ImageGrab.grab()
        filename = f"screenshot_{int(time.time())}.png"
        img.save(filename)
        speak(f"Screenshot saved as {filename}")
    except Exception as e:
        speak("Sorry, I couldn't take a screenshot.")

# Volume control functions using pycaw
def get_volume_interface():
    devices = AudioUtilities.GetSpeakers()
    interface = devices.Activate(IAudioEndpointVolume._iid_, CLSCTX_ALL, None)
    volume = cast(interface, POINTER(IAudioEndpointVolume))
    return volume

def mute_volume():
    volume = get_volume_interface()
    volume.SetMute(1, None)

def unmute_volume():
    volume = get_volume_interface()
    volume.SetMute(0, None)

def raise_volume():
    volume = get_volume_interface()
    current_volume = volume.GetMasterVolumeLevelScalar()
    new_volume = min(current_volume + 0.1, 1.0)
    volume.SetMasterVolumeLevelScalar(new_volume, None)
    speak(f"Volume raised to {int(new_volume * 100)} percent.")

def capture_photo():
    """Captures photo from webcam and saves it"""
    try:
        # Initialize camera
        cam = cv2.VideoCapture(0)
        if not cam.isOpened():
            speak("Could not access camera")
            return

        speak("Smile! Capturing photo in 3 seconds...")
        for i in range(3, 0, -1):
            speak(str(i))
            time.sleep(1)

        # Capture frame
        ret, frame = cam.read()
        cam.release()
        
        if ret:
            # Create Pictures directory if it doesn't exist
            pictures_dir = os.path.join(os.path.expanduser('~'), 'Pictures')
            os.makedirs(pictures_dir, exist_ok=True)
            
            # Save with timestamp
            photo_path = os.path.join(
                pictures_dir,
                f"photo_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.jpg"
            )
            
            # Save the image
            if cv2.imwrite(photo_path, frame):
                speak(f"Photo saved successfully")
                os.startfile(photo_path)  # Open the photo
            else:
                speak("Failed to save photo file")
        else:
            speak("Failed to capture photo")
            
    except Exception as e:
        speak("Camera error occurred")
        print(f"Camera Error Details: {str(e)}")
        
def lower_volume():
    volume = get_volume_interface()
    current_volume = volume.GetMasterVolumeLevelScalar()
    new_volume = max(current_volume - 0.1, 0.0)
    volume.SetMasterVolumeLevelScalar(new_volume, None)
    speak(f"Volume lowered to {int(new_volume * 100)} percent.")

# Main assistant loop
def run_assistant():
    wish_user()
    while True:
        wake_query = listen_for_speech()

        if 'ankit' in wake_query:
            query = listen_for_command()

            if query == "":
                continue

            if 'search' in query or 'wikipedia' in query:
                speak("Searching Wikipedia...")
                query = query.replace("wikipedia", "").replace("search", "")
                try:
                    result = wikipedia.summary(query, sentences=2)
                    speak("According to Wikipedia:")
                    speak(result)
                except:
                    speak("Sorry, I couldn't find anything.")

            elif 'open youtube' in query:
                speak("Opening YouTube...")
                webbrowser.open("https://www.youtube.com/")

            elif 'open google' in query:
                speak("Opening Google...")
                webbrowser.open("https://www.google.com/")

            elif 'close google' in query or 'close chrome' in query:
                speak("Closing Chrome browser...")
                os.system("taskkill /f /im chrome.exe")

            elif 'time' in query:
                strTime = datetime.datetime.now().strftime("%H:%M:%S")
                speak(f"The current time is {strTime}")

            elif 'date' in query:
                today = datetime.date.today().strftime("%B %d, %Y")
                speak(f"Today's date is {today}")

            elif 'day' in query:
                day = datetime.datetime.now().strftime("%A")
                speak(f"Today is {day}")

            elif 'joke' in query:
                joke = pyjokes.get_joke()
                speak(joke)

            elif 'click photo' in query or 'take photo' in query:
                capture_photo()

            elif 'fun fact' in query:
                facts = [
                    "Honey never spoils. Archaeologists have found pots of honey in ancient Egyptian tombs that are over 3000 years old and still edible.",
                    "Octopuses have three hearts and blue blood.",
                    "Bananas are berries, but strawberries are not.",
                    "A bolt of lightning contains enough energy to toast 100,000 slices of bread.",
                    "Sharks existed before trees."
                ]
                speak(random.choice(facts))

            elif 'chatgpt' in query:
                speak("Opening ChatGPT...")
                webbrowser.open("https://www.chatgpt.com/")

            elif 'weather' in query:
                speak("Showing current weather in your browser.")
                webbrowser.open("https://www.google.com/search?q=weather")

            elif 'screenshot' in query:
                take_screenshot()

            elif 'mute' in query:
                mute_volume()
                speak("Muting volume.")

            elif 'unmute' in query:
                unmute_volume()
                speak("Unmuting volume.")

            elif 'raise volume' in query or 'increase volume' in query:
                raise_volume()

            elif 'lower volume' in query or 'decrease volume' in query:
                lower_volume()

            elif 'open notepad' in query:
                speak("Opening Notepad...")
                os.system("notepad")

            elif 'open command prompt' in query or 'open cmd' in query:
                speak("Opening Command Prompt...")
                os.system("start cmd")

            elif 'lock' in query:
                speak("Locking the computer...")
                os.system("rundll32.exe user32.dll,LockWorkStation")

            elif 'shutdown' in query:
                speak("Shutting down the system...")
                os.system("shutdown /s /t 1")

            elif 'restart' in query:
                speak("Restarting the system...")
                os.system("shutdown /r /t 1")

            elif 'log out' in query or 'logout' in query:
                speak("Logging out...")
                os.system("shutdown -l")

            # ========== NEW WINDOWS OS FEATURES ==========
            elif 'create folder' in query:
                try:
                    folder_name = query.split("create folder")[-1].strip()
                    os.mkdir(folder_name)
                    speak(f"Folder '{folder_name}' created successfully.")
                except Exception as e:
                    speak(f"Failed to create folder. Error: {str(e)}")

            elif 'delete folder' in query:
                try:
                    folder_name = query.split("delete folder")[-1].strip()
                    os.rmdir(folder_name)
                    speak(f"Folder '{folder_name}' deleted successfully.")
                except Exception as e:
                    speak(f"Failed to delete folder. Error: {str(e)}")

            elif 'list files' in query:
                files = os.listdir()
                speak(f"There are {len(files)} items here. Some are: {', '.join(files[:3])}")

            elif 'open file' in query:
                try:
                    file_name = query.split("open file")[-1].strip()
                    os.startfile(file_name)
                    speak(f"Opening {file_name}")
                except Exception as e:
                    speak("File not found or invalid format.")

            elif 'delete file' in query:
                try:
                    file_name = query.split("delete file")[-1].strip()
                    os.remove(file_name)
                    speak(f"File '{file_name}' deleted.")
                except Exception as e:
                    speak("File not found or already deleted.")

            elif 'system info' in query:
                info = f"""
                System: {platform.system()}
                Version: {platform.version()}
                Machine: {platform.machine()}
                Processor: {platform.processor()}
                """
                speak("System information:" + info)

            elif 'task list' in query:
                try:
                    tasks = os.popen('tasklist').read()
                    speak("Showing running tasks in console.")
                    print(tasks)
                except:
                    speak("Couldn't fetch task list.")

            elif 'kill task' in query:
                try:
                    task_name = query.split("kill task")[-1].strip() + ".exe"
                    os.system(f"taskkill /f /im {task_name}")
                    speak(f"Terminated {task_name}")
                except:
                    speak("Failed to kill the task.")

            elif 'empty recycle bin' in query:
                try:
                    os.system('powershell.exe Clear-RecycleBin -Force')
                    speak("Recycle bin emptied.")
                except:
                    speak("Failed to empty recycle bin.")

            elif 'ip address' in query:
                try:
                    ip = os.popen('ipconfig').read()
                    speak("Your IP details are printed in console.")
                    print(ip)
                except:
                    speak("Couldn't fetch IP information.")

            elif 'sleep mode' in query:
                speak("Putting system to sleep.")
                os.system("rundll32.exe powrprof.dll,SetSuspendState 0,1,0")

            elif 'hibernate' in query:
                speak("Hibernating system.")
                os.system("shutdown /h")

            elif 'take notes' in query or 'take note' in query:
                try:
                    speak("I'm ready to take your notes. Please speak clearly.")
                    note_content = listen_for_command()
        
                    if note_content:
                        doc = Document()
                        doc.add_heading('Voice Note', 0)
                        doc.add_paragraph(note_content)
                        doc.add_paragraph(f"Recorded at: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
            
                        # Save to temporary file
                        temp_path = os.path.join(tempfile.gettempdir(), f"voice_note_{int(time.time())}.docx")
                        doc.save(temp_path)
            
                        os.startfile(temp_path)
                        speak("Note saved and opened in Word.")
                    else:
                        speak("I didn't hear any note content.")
            
                except Exception as e:
                    speak("Failed to create Word document.")                        
                    print(f"DOCX Error: {e}")

            elif 'show disk space' in query:
                try:
                    disk = os.popen('wmic logicaldisk get size,freespace,caption').read()
                    speak("Disk space details printed in console.")
                    print(disk)
                except:
                    speak("Couldn't fetch disk information.")

            elif 'open calculator' in query:
                speak("Opening Calculator.")
                os.system("calc")

            elif 'open paint' in query:
                speak("Opening Paint.")
                os.system("mspaint")

            elif 'open task manager' in query:
                speak("Opening Task Manager.")
                os.system("taskmgr")

            elif 'open control panel' in query:
                speak("Opening Control Panel.")
                os.system("control")

            elif 'open settings' in query:
                speak("Opening Windows Settings.")
                os.system("start ms-settings:")

            elif 'check updates' in query:
                speak("Checking for Windows updates.")
                os.system("start ms-settings:windowsupdate")

            elif 'open camera' in query:
                speak("Opening Camera.")
                os.system("start microsoft.windows.camera:")

          
            elif 'exit' in query or 'bye' in query or 'quit' in query:
                speak("Goodbye! Have a nice day!")
                break

            else:
                speak("Sorry, I didn't understand that.")

if __name__ == "__main__":
    run_assistant()